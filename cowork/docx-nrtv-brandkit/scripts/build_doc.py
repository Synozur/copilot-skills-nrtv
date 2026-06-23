#!/usr/bin/env python3
"""
Narrative Strategies brand-kit Word builder.

Applies the client's fonts, colors, and logo to a CLEAN, flexible layout.
It deliberately does NOT reproduce the sample template's cover page or
header/footer artwork — only the brand DNA.

Usage:
    python3 build_doc.py content.json output/Document.docx

content.json schema (all fields optional except `blocks`):
{
  "title":     "Strategic Communications Brief",
  "subtitle":  "Prepared for the Executive Team",
  "meta":      "June 23, 2026  •  Confidential",   # small line under subtitle
  "logo":      "full" | "mark" | "none",            # header logo, default "full"
  "blocks": [
    {"type": "h1", "text": "Section heading"},
    {"type": "h2", "text": "Sub heading"},
    {"type": "h3", "text": "Minor heading"},
    {"type": "p",  "text": "Body paragraph ..."},
    {"type": "bullet", "items": ["point one", "point two"]},
    {"type": "number", "items": ["step one", "step two"]},
    {"type": "quote", "text": "A pulled-out callout line."},
    {"type": "rule"},                                 # thin navy divider
    {"type": "table", "header": ["Col A","Col B"], "rows": [["1","2"],["3","4"]]}
  ]
}
"""
import sys, json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand constants (from Sample.docx) ------------------------------------
NAVY      = RGBColor(0x10, 0x3F, 0x72)
DARKNAVY  = RGBColor(0x0E, 0x28, 0x41)
SLATE     = RGBColor(0x47, 0x74, 0x99)
MUTED     = RGBColor(0x60, 0x76, 0x9C)
PALE      = "E6EBF1"   # fills are hex strings (shading)
GRAY      = "E8E8E8"
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

FONT_BODY = "Montserrat"
FONT_HEAD = "Montserrat SemiBold"
FONT_MED  = "Montserrat Medium"

HERE   = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.normpath(os.path.join(HERE, "..", "assets"))


def set_run_font(run, name, size=None, color=None, bold=False, italic=False):
    run.font.name = name
    # force east-asian + complex-script slots so the name sticks in Word
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def shade(cell, hexfill):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexfill)
    tcpr.append(sh)


def add_rule(doc, color="103F72", size_pt=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(int(size_pt*8)))
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), color)
    pbdr.append(bottom); ppr.append(pbdr)


def style_base(doc):
    """Set Normal + heading styles to the Narrative brand."""
    normal = doc.styles['Normal']
    normal.font.name = FONT_BODY; normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii','w:hAnsi','w:cs'): rpr.set(qn(a), FONT_BODY)
    specs = {
        'Heading 1': (FONT_HEAD, 14, NAVY),
        'Heading 2': (FONT_HEAD, 12, NAVY),
        'Heading 3': (FONT_MED, 10.5, NAVY),
    }
    for sname,(fn,sz,col) in specs.items():
        try:
            st = doc.styles[sname]
            st.font.name = fn; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = False
            r = st.element.get_or_add_rPr().get_or_add_rFonts()
            for a in ('w:ascii','w:hAnsi','w:cs'): r.set(qn(a), fn)
        except KeyError:
            pass


def add_header_logo(doc, which):
    if which == "none":
        return
    fn = "narrative-logo-full.png" if which != "mark" else "narrative-mark.png"
    path = os.path.join(ASSETS, fn)
    if not os.path.exists(path):
        return
    hdr = doc.sections[0].header
    p = hdr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    if which == "mark":
        run.add_picture(path, height=Inches(0.38))
    else:
        run.add_picture(path, width=Inches(1.9))


def add_title_block(doc, spec):
    if spec.get("title"):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(spec["title"]), FONT_HEAD, 24, NAVY, bold=False)
    if spec.get("subtitle"):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(spec["subtitle"]), FONT_MED, 13, SLATE)
    if spec.get("meta"):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        set_run_font(p.add_run(spec["meta"]), FONT_BODY, 9, MUTED)
    if spec.get("title") or spec.get("subtitle"):
        add_rule(doc)


def add_table(doc, header, rows):
    cols = len(header) if header else (len(rows[0]) if rows else 1)
    t = doc.add_table(rows=0, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    try: t.style = 'Table Grid'
    except KeyError: pass
    if header:
        hr = t.add_row().cells
        for i, htext in enumerate(header):
            shade(hr[i], "103F72")
            para = hr[i].paragraphs[0]
            set_run_font(para.add_run(str(htext)), FONT_HEAD, 10, WHITE, bold=False)
    for r_i, row in enumerate(rows or []):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if i >= cols: break
            if r_i % 2 == 1: shade(cells[i], PALE)
            para = cells[i].paragraphs[0]
            set_run_font(para.add_run(str(val)), FONT_BODY, 9.5, RGBColor(0x20,0x20,0x20))


def add_quote(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    ppr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'18'); left.set(qn('w:space'),'8'); left.set(qn('w:color'),'477499')
    pbdr.append(left); ppr.append(pbdr)
    set_run_font(p.add_run(text), FONT_MED, 11, MUTED, italic=True)


def build(spec, out_path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
    style_base(doc)
    add_header_logo(doc, spec.get("logo", "full"))
    add_title_block(doc, spec)

    for b in spec.get("blocks", []):
        bt = b.get("type")
        if bt == "h1": doc.add_heading(b.get("text",""), level=1)
        elif bt == "h2": doc.add_heading(b.get("text",""), level=2)
        elif bt == "h3": doc.add_heading(b.get("text",""), level=3)
        elif bt == "p": doc.add_paragraph(b.get("text",""))
        elif bt == "bullet":
            for it in b.get("items", []): doc.add_paragraph(str(it), style='List Bullet')
        elif bt == "number":
            for it in b.get("items", []): doc.add_paragraph(str(it), style='List Number')
        elif bt == "quote": add_quote(doc, b.get("text",""))
        elif bt == "rule": add_rule(doc)
        elif bt == "table": add_table(doc, b.get("header"), b.get("rows"))

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 3:
        print("usage: build_doc.py <content.json> <output.docx>", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], encoding='utf-8') as f:
        spec = json.load(f)
    path = build(spec, sys.argv[2])
    print("Wrote", path)


if __name__ == "__main__":
    main()
